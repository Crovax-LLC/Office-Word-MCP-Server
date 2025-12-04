"""
Document creation and manipulation tools for Word Document Server.
"""
import os
import json
import tempfile
from typing import Dict, List, Optional, Any
from docx import Document

from word_document_server.utils.file_utils import (
    check_file_writeable, ensure_docx_extension, create_document_copy,
    S3FileContext, check_file_exists, upload_if_s3, cleanup_temp_file
)
from word_document_server.utils.s3_utils import is_s3_uri, upload_to_s3, download_from_s3
from word_document_server.utils.document_utils import get_document_properties, extract_document_text, get_document_structure, get_document_xml, insert_header_near_text, insert_line_or_paragraph_near_text
from word_document_server.core.styles import ensure_heading_style, ensure_table_style


async def create_document(filename: str, title: Optional[str] = None, author: Optional[str] = None) -> str:
    """Create a new Word document with optional metadata.

    Args:
        filename: Name of the document to create (local path or S3 URI like s3://bucket/path/doc.docx)
        title: Optional title for the document metadata
        author: Optional author for the document metadata
    """
    filename = ensure_docx_extension(filename)

    # Handle S3 URIs
    if is_s3_uri(filename):
        try:
            # Create document in temp file, then upload to S3
            fd, local_path = tempfile.mkstemp(suffix='.docx')
            os.close(fd)

            try:
                doc = Document()

                if title:
                    doc.core_properties.title = title
                if author:
                    doc.core_properties.author = author

                ensure_heading_style(doc)
                ensure_table_style(doc)

                doc.save(local_path)

                # Upload to S3
                success, message = upload_to_s3(local_path, filename)
                if not success:
                    return f"Failed to upload to S3: {message}"

                return f"Document {filename} created successfully"
            finally:
                # Clean up temp file
                if os.path.exists(local_path):
                    os.unlink(local_path)
        except Exception as e:
            return f"Failed to create document: {str(e)}"

    # Local file handling
    is_writeable, error_message = check_file_writeable(filename)
    if not is_writeable:
        return f"Cannot create document: {error_message}"

    try:
        doc = Document()

        # Set properties if provided
        if title:
            doc.core_properties.title = title
        if author:
            doc.core_properties.author = author

        # Ensure necessary styles exist
        ensure_heading_style(doc)
        ensure_table_style(doc)

        # Save the document
        doc.save(filename)

        return f"Document {filename} created successfully"
    except Exception as e:
        return f"Failed to create document: {str(e)}"


async def get_document_info(filename: str) -> str:
    """Get information about a Word document.

    Args:
        filename: Path to the Word document (local path or S3 URI)
    """
    filename = ensure_docx_extension(filename)

    try:
        with S3FileContext(filename, read_only=True) as ctx:
            if not os.path.exists(ctx.local_path):
                return f"Document {filename} does not exist"
            properties = get_document_properties(ctx.local_path)
            return json.dumps(properties, indent=2)
    except IOError as e:
        return str(e)
    except Exception as e:
        return f"Failed to get document info: {str(e)}"


async def get_document_text(filename: str) -> str:
    """Extract all text from a Word document.

    Args:
        filename: Path to the Word document (local path or S3 URI)
    """
    filename = ensure_docx_extension(filename)

    try:
        with S3FileContext(filename, read_only=True) as ctx:
            return extract_document_text(ctx.local_path)
    except IOError as e:
        return str(e)
    except Exception as e:
        return f"Failed to extract text: {str(e)}"


async def get_document_outline(filename: str) -> str:
    """Get the structure of a Word document.

    Args:
        filename: Path to the Word document (local path or S3 URI)
    """
    filename = ensure_docx_extension(filename)

    try:
        with S3FileContext(filename, read_only=True) as ctx:
            structure = get_document_structure(ctx.local_path)
            return json.dumps(structure, indent=2)
    except IOError as e:
        return str(e)
    except Exception as e:
        return f"Failed to get document outline: {str(e)}"


async def list_available_documents(directory: str = ".") -> str:
    """List all .docx files in the specified directory.
    
    Args:
        directory: Directory to search for Word documents
    """
    try:
        if not os.path.exists(directory):
            return f"Directory {directory} does not exist"
        
        docx_files = [f for f in os.listdir(directory) if f.endswith('.docx')]
        
        if not docx_files:
            return f"No Word documents found in {directory}"
        
        result = f"Found {len(docx_files)} Word documents in {directory}:\n"
        for file in docx_files:
            file_path = os.path.join(directory, file)
            size = os.path.getsize(file_path) / 1024  # KB
            result += f"- {file} ({size:.2f} KB)\n"
        
        return result
    except Exception as e:
        return f"Failed to list documents: {str(e)}"


async def copy_document(source_filename: str, destination_filename: Optional[str] = None) -> str:
    """Create a copy of a Word document.

    Args:
        source_filename: Path to the source document (local path or S3 URI)
        destination_filename: Optional path for the copy (local path or S3 URI). If not provided, a default name will be generated.
    """
    source_filename = ensure_docx_extension(source_filename)

    if destination_filename:
        destination_filename = ensure_docx_extension(destination_filename)
    else:
        # Generate default destination
        if is_s3_uri(source_filename):
            # For S3, add _copy before extension
            base, ext = os.path.splitext(source_filename)
            destination_filename = f"{base}_copy{ext}"
        else:
            base, ext = os.path.splitext(source_filename)
            destination_filename = f"{base}_copy{ext}"

    source_is_s3 = is_s3_uri(source_filename)
    dest_is_s3 = is_s3_uri(destination_filename)

    try:
        if source_is_s3 or dest_is_s3:
            # At least one is S3, need to handle specially
            local_source = None
            temp_source = False

            try:
                # Download source if S3
                if source_is_s3:
                    success, message, local_source = download_from_s3(source_filename)
                    if not success:
                        return f"Failed to download source from S3: {message}"
                    temp_source = True
                else:
                    local_source = source_filename
                    if not os.path.exists(local_source):
                        return f"Source document {source_filename} does not exist"

                # Copy to destination
                if dest_is_s3:
                    # Upload to S3
                    success, message = upload_to_s3(local_source, destination_filename)
                    if not success:
                        return f"Failed to upload to S3: {message}"
                    return f"Document copied to {destination_filename}"
                else:
                    # Copy to local
                    import shutil
                    shutil.copy2(local_source, destination_filename)
                    return f"Document copied to {destination_filename}"
            finally:
                # Clean up temp file
                if temp_source and local_source and os.path.exists(local_source):
                    os.unlink(local_source)
        else:
            # Both are local
            success, message, new_path = create_document_copy(source_filename, destination_filename)
            if success:
                return message
            else:
                return f"Failed to copy document: {message}"
    except Exception as e:
        return f"Failed to copy document: {str(e)}"


async def merge_documents(target_filename: str, source_filenames: List[str], add_page_breaks: bool = True) -> str:
    """Merge multiple Word documents into a single document.

    Args:
        target_filename: Path to the target document (local path or S3 URI)
        source_filenames: List of paths to source documents to merge (local paths or S3 URIs)
        add_page_breaks: If True, add page breaks between documents
    """
    from word_document_server.core.tables import copy_table

    target_filename = ensure_docx_extension(target_filename)
    target_is_s3 = is_s3_uri(target_filename)

    # Check if target file is writeable (only for local files)
    if not target_is_s3:
        is_writeable, error_message = check_file_writeable(target_filename)
        if not is_writeable:
            return f"Cannot create target document: {error_message}"

    # Track temp files for cleanup
    temp_files = []
    local_target = None

    try:
        # Resolve all source files (download from S3 if needed)
        local_sources = []
        for filename in source_filenames:
            doc_filename = ensure_docx_extension(filename)

            if is_s3_uri(doc_filename):
                success, message, local_path = download_from_s3(doc_filename)
                if not success:
                    return f"Failed to download {doc_filename} from S3: {message}"
                local_sources.append(local_path)
                temp_files.append(local_path)
            else:
                if not os.path.exists(doc_filename):
                    return f"Source document {doc_filename} does not exist"
                local_sources.append(doc_filename)

        # Create target file (temp if S3)
        if target_is_s3:
            fd, local_target = tempfile.mkstemp(suffix='.docx')
            os.close(fd)
            temp_files.append(local_target)
        else:
            local_target = target_filename

        # Create a new document for the merged result
        target_doc = Document()

        # Process each source document
        for i, local_source in enumerate(local_sources):
            source_doc = Document(local_source)

            # Add page break between documents (except before the first one)
            if add_page_breaks and i > 0:
                target_doc.add_page_break()

            # Copy all paragraphs
            for paragraph in source_doc.paragraphs:
                # Create a new paragraph with the same text and style
                new_paragraph = target_doc.add_paragraph(paragraph.text)
                new_paragraph.style = target_doc.styles['Normal']  # Default style

                # Try to match the style if possible
                try:
                    if paragraph.style and paragraph.style.name in target_doc.styles:
                        new_paragraph.style = target_doc.styles[paragraph.style.name]
                except:
                    pass

                # Copy run formatting
                for j, run in enumerate(paragraph.runs):
                    if j < len(new_paragraph.runs):
                        new_run = new_paragraph.runs[j]
                        # Copy basic formatting
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        # Font size if specified
                        if run.font.size:
                            new_run.font.size = run.font.size

            # Copy all tables
            for table in source_doc.tables:
                copy_table(table, target_doc)

        # Save the merged document
        target_doc.save(local_target)

        # Upload to S3 if needed
        if target_is_s3:
            success, message = upload_to_s3(local_target, target_filename)
            if not success:
                return f"Failed to upload merged document to S3: {message}"

        return f"Successfully merged {len(source_filenames)} documents into {target_filename}"
    except Exception as e:
        return f"Failed to merge documents: {str(e)}"
    finally:
        # Clean up temp files
        for temp_file in temp_files:
            if temp_file and os.path.exists(temp_file):
                try:
                    os.unlink(temp_file)
                except:
                    pass


async def get_document_xml_tool(filename: str) -> str:
    """Get the raw XML structure of a Word document.

    Args:
        filename: Path to the Word document (local path or S3 URI)
    """
    filename = ensure_docx_extension(filename)

    try:
        with S3FileContext(filename, read_only=True) as ctx:
            return get_document_xml(ctx.local_path)
    except IOError as e:
        return str(e)
    except Exception as e:
        return f"Failed to get document XML: {str(e)}"
